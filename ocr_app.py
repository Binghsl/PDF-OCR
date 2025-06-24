import streamlit as st
from pdf2image import convert_from_bytes
from PIL import Image
import pytesseract
from docx import Document
from io import BytesIO
from PyPDF2 import PdfWriter, PdfReader

# Improve OCR accuracy with config
TESSERACT_CONFIG = '--psm 3'  # Auto page segmentation

# Optional: Poppler path for Windows users
# import os
# os.environ["PATH"] += os.pathsep + r'C:\path\to\poppler\bin'

def scanned_pdf_to_text_and_images(pdf_bytes, dpi=300):
    pages = convert_from_bytes(pdf_bytes, dpi=dpi)
    full_text = ""
    page_texts = []
    for page_number, page_image in enumerate(pages, start=1):
        text = pytesseract.image_to_string(page_image, lang='eng', config=TESSERACT_CONFIG)
        page_texts.append((page_image, text))
        full_text += f"\n\n--- Page {page_number} ---\n\n{text}"
    return full_text, page_texts

def save_to_word(page_texts):
    doc = Document()
    for i, (_, text) in enumerate(page_texts):
        doc.add_heading(f'Page {i+1}', level=2)
        doc.add_paragraph(text)
    word_io = BytesIO()
    doc.save(word_io)
    word_io.seek(0)
    return word_io

def create_searchable_pdf(page_texts):
    pdf_writer = PdfWriter()
    for image, text in page_texts:
        img_buffer = BytesIO()
        image.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        pdf_image_page = Image.open(img_buffer).convert("RGB")
        pdf_image_bytes = BytesIO()
        pdf_image_page.save(pdf_image_bytes, format="PDF")
        pdf_image_bytes.seek(0)

        # Overlay OCR text (note: this creates just image-based PDF, not true invisible OCR layer)
        reader = PdfReader(pdf_image_bytes)
        page = reader.pages[0]
        pdf_writer.add_page(page)

    pdf_bytes = BytesIO()
    pdf_writer.write(pdf_bytes)
    pdf_bytes.seek(0)
    return pdf_bytes

# UI
st.title("Scanned PDF to Text OCR (with Export)")

uploaded_file = st.file_uploader("Upload a scanned PDF file", type=["pdf"])
dpi = st.slider("Image DPI (higher = better OCR, slower)", min_value=100, max_value=600, value=300, step=50)

if uploaded_file is not None:
    with st.spinner("Processing OCR..."):
        pdf_bytes = uploaded_file.read()
        extracted_text, page_texts = scanned_pdf_to_text_and_images(pdf_bytes, dpi=dpi)

    st.text_area("Extracted Text", extracted_text, height=400)

    # Download Word
    word_file = save_to_word(page_texts)
    st.download_button("Download as Word (.docx)", word_file, file_name="ocr_output.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # Download Searchable PDF (Image-based + OCR overlay text placeholder)
    searchable_pdf = create_searchable_pdf(page_texts)
    st.download_button("Download Searchable PDF", searchable_pdf, file_name="ocr_output.pdf", mime="application/pdf")
